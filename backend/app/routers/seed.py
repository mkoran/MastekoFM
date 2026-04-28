"""Seed endpoints — populate DEV with canonical scenarios from seed/* files.

Sprint G1: rewritten under the new Workspace + per-folder + versioned-filename
layout. Each entity gets its own Drive folder. All files use {code}_v001.xlsx.

  Workspaces/{ws_code}/
    Models/{model_code}/{model_code}_v001.xlsx
    OutputTemplates/{tpl_code}/{tpl_code}_v001.xlsx
    Projects/{project_code}/
      AssumptionPacks/{pack_code}/{pack_code}_v001.xlsx
      Runs/{ts}_{pack}_{tpl}/{pack}_{tpl}_v001.xlsx (created at run time)

The seed auto-creates a "Personal" workspace for the calling user if they
don't have one. Idempotent: re-running returns existing ids.
"""
import io
import logging
from datetime import UTC, datetime
from pathlib import Path
from typing import Annotated, Any

from docx import Document  # type: ignore[import-untyped]
from fastapi import APIRouter, Depends, HTTPException, Request

from backend.app.config import get_firestore_client, settings
from backend.app.middleware.auth import get_current_user
from backend.app.services import drive_service, excel_template_engine, storage_service

logger = logging.getLogger(__name__)
router = APIRouter(tags=["seed"])

CurrentUser = Annotated[dict[str, Any], Depends(get_current_user)]
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

REPO_ROOT = Path(__file__).resolve().parents[3]


def _get_or_create_personal_workspace(uid: str, email: str, root: str, user_token: str) -> tuple[str, str]:
    """Returns (workspace_id, workspace_code). Auto-creates 'Personal' if none."""
    db = get_firestore_client()
    coll = db.collection(f"{_prefix()}workspaces")
    for doc in coll.where("members", "array_contains", uid).limit(1).stream():
        d = doc.to_dict()
        return doc.id, d.get("code_name", "personal")
    # Create new
    code = storage_service.safe_name(f"personal-{uid[:6]}", fallback="personal")
    folder_id = drive_service.ensure_workspace_folders(root, code, user_access_token=user_token)["workspace"]
    now = datetime.now(UTC)
    new_ref = coll.document()
    new_ref.set({
        "name": "Personal",
        "code_name": code,
        "description": f"Auto-created via seed for {email}",
        "members": [uid],
        "drive_folder_id": folder_id,
        "archived": False,
        "created_by": uid,
        "created_by_email": email,
        "created_at": now,
        "updated_at": now,
    })
    return new_ref.id, code


def _prefix() -> str:
    return settings.firestore_collection_prefix


def _drive_root_id() -> str:
    doc = get_firestore_client().collection(f"{_prefix()}settings").document("app").get()
    if doc.exists:
        return (doc.to_dict() or {}).get("drive_root_folder_id") or settings.drive_root_folder_id
    return settings.drive_root_folder_id


def _find_by_code(collection: str, code: str) -> tuple[str | None, dict[str, Any] | None]:
    db = get_firestore_client()
    for doc in db.collection(collection).stream():
        data = doc.to_dict()
        if data.get("code_name") == code:
            return doc.id, data
    return None, None


def _find_pack_by_code(project_id: str, code: str) -> tuple[str | None, dict[str, Any] | None]:
    db = get_firestore_client()
    ref = db.collection(f"{_prefix()}projects").document(project_id).collection("assumption_packs")
    for doc in ref.stream():
        data = doc.to_dict()
        if data.get("code_name") == code:
            return doc.id, data
    return None, None


def _build_helloworld_narrative_docx_bytes() -> bytes:
    """Sprint D-2 — generate a Hello World narrative DOCX template.

    Marc opens this Google Doc in his browser to edit (true WYSIWYG). At
    Run time, MastekoFM exports the Doc as DOCX, fills the Jinja2
    placeholders with Model output values, and renders a PDF.

    Cell references match the canonical Hello World fixture
    (``scripts/build_helloworld_seed.py``):

        I_Numbers.B1 = a (5),  I_Numbers.B2 = b (7)
        O_Results.B1 = sum (12), O_Results.B2 = product (35)
    """
    doc = Document()
    doc.add_heading("Hello World Report", level=0)
    doc.add_paragraph(
        "Generated for project {{ run.project_name }} on {{ run.started_at }}."
    )
    doc.add_heading("Inputs", level=1)
    doc.add_paragraph("a = {{ model.I_Numbers.B1 }}")
    doc.add_paragraph("b = {{ model.I_Numbers.B2 }}")
    doc.add_heading("Calculation results", level=1)
    p = doc.add_paragraph()
    p.add_run("Sum: ").bold = True
    p.add_run("{{ model.O_Results.B1 }}")
    p = doc.add_paragraph()
    p.add_run("Product: ").bold = True
    p.add_run("{{ model.O_Results.B2 }}")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "The sum of {{ model.I_Numbers.B1 }} and {{ model.I_Numbers.B2 }} is "
        "{{ model.O_Results.B1 }}; their product is {{ model.O_Results.B2 }}."
    )
    doc.add_paragraph(
        "Edit this Google Doc in your browser — change the wording, add your logo, "
        "drop in tables. Every Run will re-fill the {{ }} placeholders and produce "
        "a fresh PDF."
    )
    doc.add_paragraph(
        "Available placeholders: {{ model.<TabName>.<Cell> }} for any I_* or O_* "
        "cell on the Model, plus run-level fields: {{ run.id }}, {{ run.project_name }}, "
        "{{ run.model_name }}, {{ run.pack_name }}, {{ run.started_at }}."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed_one(
    *,
    seed_dir: Path,
    model_filename: str,
    pack_filename: str,
    output_template_filename: str,
    model_name: str,
    model_code: str,
    model_description: str,
    pack_name: str,
    pack_code: str,
    pack_description: str,
    template_name: str,
    template_code: str,
    template_description: str,
    project_name: str,
    project_code: str,
    project_description: str,
    user_token: str,
    current_user: dict,
    pdf_export_xlsx: bool = False,  # Sprint D-1: opt the seed's OutputTemplate into PDF
    narrative_docx_factory: "Any" = None,  # Sprint D-2: callable() -> docx bytes
    narrative_template_name: str | None = None,  # Sprint D-2: filename for the Google Doc
) -> dict[str, Any]:
    """Generic seeder used by both Hello World and Campus Adele endpoints."""
    root = _drive_root_id()
    if not root:
        raise HTTPException(status_code=400, detail="No Drive root folder configured.")

    db = get_firestore_client()
    now = datetime.now(UTC)
    result: dict[str, Any] = {"created": [], "existing": []}

    model_path = seed_dir / model_filename
    pack_path = seed_dir / pack_filename
    tpl_path = seed_dir / output_template_filename
    for p in (model_path, pack_path, tpl_path):
        if not p.exists():
            raise HTTPException(status_code=500, detail=f"Seed file missing: {p}")

    # ── 0. Workspace (Sprint G1) ────────────────────────────────────────────
    ws_id, ws_code = _get_or_create_personal_workspace(
        current_user["uid"], current_user.get("email", ""), root, user_token,
    )
    ws_folders = drive_service.ensure_workspace_folders(root, ws_code, user_access_token=user_token)

    # ── 1. Model (Drive-backed; Sprint G1) ──────────────────────────────────
    model_id, model_data = _find_by_code(f"{_prefix()}models", model_code)
    if model_id:
        result["existing"].append(f"model={model_id}")
    else:
        content = model_path.read_bytes()
        classes = excel_template_engine.classify_bytes(content)
        model_folder_id = drive_service.ensure_model_folder(
            ws_folders["models"], model_code, user_access_token=user_token,
        )
        filename = drive_service.versioned_filename(model_code, 1, ext="xlsx")
        drive_id = drive_service.upload_file(
            model_folder_id, filename, content, XLSX_MIME, user_access_token=user_token,
        )
        model_ref = db.collection(f"{_prefix()}models").document()
        model_data = {
            "name": model_name,
            "code_name": model_code,
            "description": model_description,
            "workspace_id": ws_id,
            "version": 1,
            "input_tabs": classes["input_tabs"],
            "output_tabs": classes["output_tabs"],
            "calc_tabs": classes["calc_tabs"],
            "storage_path": None,
            "drive_folder_id": model_folder_id,
            "drive_file_id": drive_id,
            "size_bytes": len(content),
            "uploaded_by": current_user["uid"],
            "uploaded_by_email": current_user.get("email", ""),
            "created_by_email": current_user.get("email", ""),
            "created_at": now,
            "updated_at": now,
        }
        model_ref.set(model_data)
        model_id = model_ref.id
        result["created"].append(f"model={model_id}")

    # ── 2. OutputTemplate (Drive-backed; per-template folder; Sprint G1) ────
    tpl_id, tpl_data = _find_by_code(f"{_prefix()}output_templates", template_code)
    if tpl_id:
        # Idempotent flag-patch: keep an existing OutputTemplate's behaviour
        # in sync with what the seed declares today. Without this, an old
        # Hello World seed (pre-Sprint D-1) would never produce PDFs unless
        # Marc wiped + reseeded.
        patch: dict[str, Any] = {}
        if (tpl_data or {}).get("pdf_export_xlsx") != pdf_export_xlsx:
            patch["pdf_export_xlsx"] = pdf_export_xlsx
        # Sprint D-2: if a narrative template is requested but the existing
        # OutputTemplate has none, create it now and attach the new file id.
        # If one already exists, leave it alone — Marc may have edited it.
        if (
            narrative_docx_factory is not None
            and narrative_template_name
            and not (tpl_data or {}).get("google_doc_template_drive_file_id")
        ):
            try:
                tpl_folder_id_existing = (tpl_data or {}).get("drive_folder_id")
                if tpl_folder_id_existing:
                    new_doc_id = drive_service.upload_docx_as_google_doc(
                        tpl_folder_id_existing,
                        narrative_template_name,
                        narrative_docx_factory(),
                        user_access_token=user_token,
                    )
                    if new_doc_id:
                        patch["google_doc_template_drive_file_id"] = new_doc_id
            except Exception:  # noqa: BLE001 — best-effort
                logger.exception(
                    "Sprint D-2: narrative template upload failed for tpl %s", tpl_id,
                )
        if patch:
            patch["updated_at"] = now
            db.collection(f"{_prefix()}output_templates").document(tpl_id).update(patch)
            result["updated"] = result.get("updated", [])
            result["updated"].append(
                f"output_template={tpl_id}: " + ", ".join(f"{k}={v}" for k, v in patch.items() if k != "updated_at")
            )
        result["existing"].append(f"output_template={tpl_id}")
    else:
        content = tpl_path.read_bytes()
        classes = excel_template_engine.classify_bytes(content)
        tpl_folder_id = drive_service.ensure_output_template_folder(
            ws_folders["output_templates"], template_code, user_access_token=user_token,
        )
        filename = drive_service.versioned_filename(template_code, 1, ext="xlsx")
        drive_id = drive_service.upload_file(
            tpl_folder_id, filename, content, XLSX_MIME, user_access_token=user_token,
        )
        tpl_ref = db.collection(f"{_prefix()}output_templates").document()
        tpl_data = {
            "name": template_name,
            "code_name": template_code,
            "description": template_description,
            "workspace_id": ws_id,
            "format": "xlsx",
            "version": 1,
            "storage_kind": "drive_xlsx",
            "storage_path": None,
            "drive_folder_id": tpl_folder_id,
            "drive_file_id": drive_id,
            "m_tabs": classes["m_tabs"],
            "output_tabs": classes["output_tabs"],
            "calc_tabs": classes["calc_tabs"],
            "size_bytes": len(content),
            "uploaded_by": current_user["uid"],
            "uploaded_by_email": current_user.get("email", ""),
            "pdf_export_xlsx": pdf_export_xlsx,  # Sprint D-1
            "created_at": now,
            "updated_at": now,
        }
        # Sprint D-2: optionally seed a narrative Google Doc template alongside.
        if narrative_docx_factory is not None and narrative_template_name:
            try:
                narrative_doc_id = drive_service.upload_docx_as_google_doc(
                    tpl_folder_id,
                    narrative_template_name,
                    narrative_docx_factory(),
                    user_access_token=user_token,
                )
                tpl_data["google_doc_template_drive_file_id"] = narrative_doc_id
                logger.info(
                    "Sprint D-2: created narrative Google Doc %s for %s",
                    narrative_doc_id, template_code,
                )
            except Exception:  # noqa: BLE001 — best-effort; xlsx still ships
                logger.exception(
                    "Sprint D-2: narrative Google Doc seed failed for %s", template_code,
                )
        tpl_ref.set(tpl_data)
        tpl_id = tpl_ref.id
        result["created"].append(f"output_template={tpl_id}")

    # ── 3. Project (Sprint G1: workspace_id + per-project folder) ───────────
    proj_id, proj_data = _find_by_code(f"{_prefix()}projects", project_code)
    if proj_id:
        result["existing"].append(f"project={proj_id}")
    else:
        proj_folders = drive_service.ensure_project_folder_v2(
            ws_folders["projects"], project_code, user_access_token=user_token,
        )
        proj_ref = db.collection(f"{_prefix()}projects").document()
        proj_data = {
            "name": project_name,
            "code_name": project_code,
            "description": project_description,
            "workspace_id": ws_id,
            "workspace_name": ws_code,
            "default_model_id": model_id,
            "default_model_name": model_name,
            "default_model_version": 1,
            "status": "active",
            "archived": False,
            # Sprint G1: store the new per-project folder ids
            "drive_folders": {
                "project": proj_folders["project"],
                "packs": proj_folders["packs"],
                "runs": proj_folders["runs"],
                # Legacy keys for back-compat with old code paths
                "inputs": proj_folders["packs"],
                "outputs": proj_folders["runs"],
            },
            "created_by": current_user["uid"],
            "created_by_email": current_user.get("email", ""),
            "created_at": now,
            "updated_at": now,
        }
        proj_ref.set(proj_data)
        proj_id = proj_ref.id
        result["created"].append(f"project={proj_id}")

    # ── 4. AssumptionPack (per-pack folder + versioned filename; Sprint G1) ─
    pack_id, pack_data = _find_pack_by_code(proj_id, pack_code)
    if pack_id:
        result["existing"].append(f"assumption_pack={pack_id}")
    else:
        content = pack_path.read_bytes()
        classes = excel_template_engine.classify_bytes(content)
        # Re-derive proj packs folder id from project doc
        proj_doc = db.collection(f"{_prefix()}projects").document(proj_id).get()
        packs_folder = (proj_doc.to_dict() or {}).get("drive_folders", {}).get("packs")
        if not packs_folder:
            # Defensive fallback
            proj_folders = drive_service.ensure_project_folder_v2(
                ws_folders["projects"], project_code, user_access_token=user_token,
            )
            packs_folder = proj_folders["packs"]
        pack_folder_id = drive_service.ensure_pack_folder(
            packs_folder, pack_code, user_access_token=user_token,
        )
        filename = drive_service.versioned_filename(pack_code, 1, ext="xlsx")
        drive_id = drive_service.upload_file(
            pack_folder_id, filename, content, XLSX_MIME, user_access_token=user_token,
        )
        pack_ref = (
            db.collection(f"{_prefix()}projects")
            .document(proj_id)
            .collection("assumption_packs")
            .document()
        )
        pack_data = {
            "name": pack_name,
            "code_name": pack_code,
            "description": pack_description,
            "project_id": proj_id,
            "pack_number": 1,                # Sprint G3: seeded pack is AP01
            "status": "active",
            "archived": False,
            "storage_kind": "drive_xlsx",
            "storage_path": None,
            "drive_folder_id": pack_folder_id,
            "drive_file_id": drive_id,
            "size_bytes": len(content),
            "version": 1,
            "input_tabs": classes["input_tabs"],
            "last_run": None,
            "created_by": current_user["uid"],
            "created_by_email": current_user.get("email", ""),
            "created_at": now,
            "updated_at": now,
        }
        pack_ref.set(pack_data)
        pack_id = pack_ref.id
        result["created"].append(f"assumption_pack={pack_id}")

    return {
        "workspace_id": ws_id,
        "workspace_code": ws_code,
        "model_id": model_id,
        "output_template_id": tpl_id,
        "project_id": proj_id,
        "assumption_pack_id": pack_id,
        **result,
    }


# ── HELLO WORLD ──────────────────────────────────────────────────────────────


def _seed_one_with_clear_errors(**kwargs):
    """Wraps _seed_one to convert internal exceptions into HTTPException with
    the original message in the body. Without this, FastAPI's default 500
    swallows the upstream error, making the frontend show only "API error: 500"
    and Cloud Logging show an empty ERROR record.

    Always logs the exception with full traceback so we can debug from logs
    too, not only from the response body.
    """
    log = logger
    try:
        return _seed_one(**kwargs)
    except HTTPException:
        raise
    except RuntimeError as exc:
        log.exception("Seed RuntimeError")
        msg = str(exc)
        if "storageQuotaExceeded" in msg or "storage quota" in msg.lower():
            raise HTTPException(
                status_code=403,
                detail=(
                    "Drive storageQuotaExceeded — the signed-in account cannot "
                    "create files at the Drive location used. Most common cause: "
                    "service-account context (e.g. CI), since SAs have no Drive "
                    "storage quota. Real users normally have quota. Fix: sign "
                    "in as a real user, or migrate the MastekoFM Drive root "
                    "into a Shared Drive. Original: " + msg
                ),
            ) from exc
        raise HTTPException(status_code=500, detail=f"Seed failed: {msg}") from exc
    except Exception as exc:  # noqa: BLE001
        log.exception("Seed unexpected exception")
        raise HTTPException(status_code=500, detail=f"Seed failed: {type(exc).__name__}: {exc}") from exc


@router.post("/api/seed/helloworld")
async def seed_helloworld(current_user: CurrentUser, request: Request):
    """Idempotent: uploads 3 Hello World seed files + creates a Project."""
    user_token = request.headers.get("X-MFM-Drive-Token")
    if not user_token:
        raise HTTPException(
            status_code=400, detail="Hello World seed requires X-MFM-Drive-Token header.",
        )
    return _seed_one_with_clear_errors(
        seed_dir=REPO_ROOT / "seed" / "helloworld",
        model_filename="helloworld_model.xlsx",
        pack_filename="helloworld_inputs.xlsx",
        output_template_filename="helloworld_report.xlsx",
        model_name="Hello World Model",
        model_code="helloworld_model",
        model_description="Tiny Hello World Model: I_Numbers (a, b) → O_Results (sum, product).",
        pack_name="Hello World Inputs",
        pack_code="helloworld_inputs",
        pack_description="Hello World inputs: a=5, b=7. Expected outputs: sum=12, product=35, total=47.",
        template_name="Hello World Report",
        template_code="helloworld_report",
        template_description="Hello World Report: pulls Model O_Results into M_Results, displays in O_Report.",
        project_name="Hello World",
        project_code="helloworld",
        project_description="Tiny Hello World Project — verifies three-way composition end-to-end.",
        user_token=user_token,
        current_user=current_user,
        pdf_export_xlsx=True,  # Sprint D-1: showcase PDF artifact alongside xlsx
        narrative_docx_factory=_build_helloworld_narrative_docx_bytes,  # Sprint D-2
        narrative_template_name="Hello World Narrative",                # → Google Doc
    )


# ── CAMPUS ADELE ─────────────────────────────────────────────────────────────


@router.post("/api/seed/campus-adele")
async def seed_campus_adele(current_user: CurrentUser, request: Request):
    """Idempotent: uploads Campus Adele Model + Base AssumptionPack + Investor Summary OutputTemplate."""
    user_token = request.headers.get("X-MFM-Drive-Token")
    if not user_token:
        raise HTTPException(
            status_code=400, detail="Campus Adele seed requires X-MFM-Drive-Token header.",
        )
    return _seed_one_with_clear_errors(
        seed_dir=REPO_ROOT / "seed" / "campus_adele",
        model_filename="campus_adele_model.xlsx",
        pack_filename="campus_adele_base_pack.xlsx",
        output_template_filename="campus_adele_summary.xlsx",
        model_name="Campus Adele (Construction-to-Perm)",
        model_code="campus_adele_model",
        model_description="15-tab construction-to-permanent financing model. 5 I_ input tabs, 1 O_ output tab.",
        pack_name="Base Case",
        pack_code="campus_adele_base",
        pack_description="Base Case AssumptionPack — 64 units, 170 budget items, all I_ tabs populated.",
        template_name="Investor Summary v1",
        template_code="campus_adele_summary",
        template_description="Minimal investor summary — surfaces Annual Summary cells from the Model.",
        project_name="Campus Adele",
        project_code="campus_adele",
        project_description="Real construction-to-perm financing project — Quebec, 64 units.",
        user_token=user_token,
        current_user=current_user,
    )
